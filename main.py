import math
import sys
import os

from PyQt6.QtGui import QIntValidator, QDoubleValidator, QRegularExpressionValidator
from dotenv import load_dotenv
import xlwt

from PyQt6 import QtWidgets, QtGui, QtCore
from PyQt6.QtCore import pyqtSignal, QSortFilterProxyModel, QRegularExpression
from PyQt6.QtSql import QSqlDatabase, QSqlQuery
from PyQt6.QtWidgets import QApplication, QWidget, QMessageBox, QFileDialog
from UI import py_ui
from utils.update_table_views import update_table_views
import Models
from PyQt6.QtCore import Qt
from docx import Document

# from Widgets.edit_record_nir import EditRecordNir
import Widgets

load_dotenv()


def connect_db(db_name=None):
    try:
        db = QSqlDatabase.addDatabase("QSQLITE")
        db.setDatabaseName(db_name)
    except:
        msgBox = QMessageBox()
        msgBox.setText("Ошибка подключения к БД")
        msgBox.exec()
        sys.exit()
    return db


class App(QWidget):
    def __init__(self):
        QWidget.__init__(self)
        self.w = QtWidgets.QMainWindow()
        self.w_root = py_ui.form.Ui_MainWindow()
        self.w_root.setupUi(self.w)
        self.w.setWindowTitle("Управление организацией НИР")

        self.db = connect_db(os.getenv('DB_URL'))

        self.nir_model = Models.NirModel(self.db)
        self.vuz_model = Models.VuzModel(self.db)
        self.fin_vuz_model = Models.FinanceVuzModel(self.db)
        self.anylyze_by_vuz = Models.AnalyzeByVuz(self.db, self.w_root)
        self.anylyze_by_grnti = Models.AnalyzeByGRNTI(self.db, self.w_root)
        self.anylyze_by_char = Models.AnalyzeByCHAR(self.db, self.w_root)
        self.finance_order = Models.FinanceOrder(self.db)

        self.w_root.tableView.setModel(self.nir_model)
        self.w_root.tableView_2.setModel(self.fin_vuz_model)
        self.w_root.tableView_3.setModel(self.vuz_model)
        self.w_root.tableView_4.setModel(self.anylyze_by_vuz)
        self.w_root.tableView_11.setModel(self.anylyze_by_grnti)
        self.w_root.tableView_12.setModel(self.anylyze_by_char)
        self.w_root.tableView_17.setModel(self.finance_order)

        # self.w_root.tableView.sortByColumn(3, Qt.SortOrder.AscendingOrder)
        self.onNirHeaderClicked(3)
        self.w_root.tableView_2.sortByColumn(0, Qt.SortOrder.AscendingOrder)
        self.w_root.tableView_3.sortByColumn(3, Qt.SortOrder.AscendingOrder)
        self.w_root.tableView.horizontalHeader().sectionClicked.connect(self.onNirHeaderClicked)

        self.w_root.action_1.triggered.connect(lambda: self.w_root.stackedWidget.setCurrentWidget(self.w_root.page_1))
        self.w_root.action_2.triggered.connect(lambda: self.w_root.stackedWidget.setCurrentWidget(self.w_root.page_2))
        self.w_root.action_3.triggered.connect(lambda: self.w_root.stackedWidget.setCurrentWidget(self.w_root.page_3))

        update_table_views(self.w_root.tableView, self.w_root.tableView_2, self.w_root.tableView_3,
                           self.w_root.tableView_4, self.w_root.tableView_11, self.w_root.tableView_12,
                           self.w_root.tableView_17)

        self.delete_accept_form = Widgets.DeleteAccept()
        self.edit_record_nir_form = Widgets.EditRecordNir()
        self.filter_nir_form = Widgets.FilterNir(self.vuz_model, self.nir_model)

        # Анализ
        self.w_root.action_7.triggered.connect(self.open_analyze)
        self.w_root.action_7.triggered.connect(lambda: self.w_root.stackedWidget.setCurrentWidget(self.w_root.page_4))

        # Финансирование
        self.w_root.action_8.triggered.connect(lambda: self.w_root.stackedWidget.setCurrentWidget(self.w_root.page_12))
        self.w_root.action_8.triggered.connect(self.open_finance)

        self.w_root.pushButton.clicked.connect(self.open_edit_record_nir)
        self.w_root.pushButton_3.clicked.connect(self.delete_records_nir)
        self.w_root.pushButton_2.clicked.connect(self.add_record_nir_form)
        self.w_root.pushButton_4.clicked.connect(self.open_filter_form)
        self.w_root.pushButton_5.clicked.connect(self.sort_by_codvuz_rnw)
        self.w_root.pushButton_6.clicked.connect(self.reset_filters)
        self.w_root.pushButton_13.clicked.connect(lambda: self.save_to_doc(self.anylyze_by_vuz, 'по вузам'))
        self.w_root.pushButton_15.clicked.connect(lambda: self.save_to_doc(self.anylyze_by_grnti, 'по ГРНТИ'))
        self.w_root.pushButton_14.clicked.connect(lambda: self.save_to_doc(self.anylyze_by_char, 'по характеру НИР'))
        self.w_root.pushButton_19.clicked.connect(lambda: self.finance_order.update(self.w_root.lineEdit_6.text()))
        self.w_root.pushButton_20.clicked.connect(self.accept_finance_order)
        self.w_root.pushButton_21.clicked.connect(self.save_to_doc_finance)


        self.w_root.lineEdit_4.textEdited.connect(self.sum_changed)
        self.w_root.lineEdit_6.textEdited.connect(self.perc_changed)

        regex = QRegularExpression("[1-9]\\d{0,20}")
        validator = QRegularExpressionValidator(regex)
        self.w_root.lineEdit_4.setValidator(validator)
        # self.w_root.lineEdit_4.setInputMask("9")
        self.w_root.lineEdit_6.setValidator(QDoubleValidator(0, float('inf'), 3))

        self.w.show()

    def save_to_doc_finance(self):
        records = []
        document = Document()
        document.add_heading(f"Распоряжение о финансировании Вузов", 0)
        query = QSqlQuery(self.finance_order.query().lastQuery())
        while query.next():
            row = []
            for i in range(self.finance_order.columnCount()):
                row.append(query.value(i))
            records.append(row)
        table = document.add_table(rows=1, cols=self.finance_order.columnCount(), style="Table Grid")
        row = table.rows[0].cells
        for i in range(self.finance_order.columnCount()):
            row[i].text = self.finance_order.headerData(i, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole)

        for row in records:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                cells[i].text = str(item)


        query = QSqlQuery(f""" SELECT SUM(z3) FROM ({self.finance_order.query().lastQuery()})""")
        while query.next():
            summa = query.value(0)


        cells = table.add_row().cells
        cells[0].text = "Итого"
        cells[1].text = str(summa)

        filename, _ = QFileDialog.getSaveFileName(
            None, 'Save Doc', os.getcwd())
        if filename:
            document.save(filename)

    def accept_finance_order(self):
        print(self.finance_order.query().lastQuery())
        QSqlQuery(f"""UPDATE Tp_fv
                        SET z18 = table2.z3
                        FROM ({self.finance_order.query().lastQuery()}) AS table2
                        WHERE Tp_fv.z2 = table2.z2""")
        self.fin_vuz_model.update()
        self.set_reference()

    def sum_changed(self, text):
        if text:
            try:
                percent = int(text) / int(self.w_root.lineEdit.text())
                self.w_root.lineEdit_6.setText(str(percent * 100))
            except:
                pass

    def perc_changed(self, text):
        if text:
            try:
                sum = float(text) * 0.01 * int(self.w_root.lineEdit.text())
                self.w_root.lineEdit_4.setText(str(int(sum)))
            except:
                pass

    def open_finance(self):
        self.set_reference()

    def set_reference(self):
        sum_plan_fin, sum_fact_fin, percent = self.get_sum_plan_fact_percent()

        self.w_root.lineEdit.setText(str(sum_plan_fin))
        self.w_root.lineEdit_2.setText(str(sum_fact_fin))
        self.w_root.lineEdit_5.setText(str(percent * 100))

    def get_sum_plan_fact_percent(self):
        query = QSqlQuery(f"""
                                SELECT SUM(f18) FROM ({self.nir_model.query().lastQuery()})
                                """)
        while query.next():
            sum_plan_fin = query.value(0)

        query = QSqlQuery(f"""
                                    SELECT SUM(z18) FROM TP_fv
                                    """)

        while query.next():
            sum_fact_fin = query.value(0)

        percent = sum_fact_fin / sum_plan_fin

        return (sum_plan_fin, sum_fact_fin, percent)

    def save_to_doc(self, model: Models, text):
        records = []
        document = Document()
        document.add_heading(f"Распределение НИР {text}", 0)
        document.add_paragraph(f"Фильтры: \n {self.w_root.textBrowser.toPlainText()}")
        query = QSqlQuery(model.query().lastQuery())
        while query.next():
            row = []
            for i in range(model.columnCount()):
                row.append(query.value(i))
            records.append(row)
        table = document.add_table(rows=1, cols=model.columnCount(), style="Table Grid")
        row = table.rows[0].cells
        for i in range(model.columnCount()):
            row[i].text = model.headerData(i, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole)

        for row in records:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                cells[i].text = str(item)
        # cells = table.add_row().cells
        # cells[0].text = 'Итого'
        # cells[1].text =str(self.w_root.label_9.text())
        # cells[2].text =str(self.w_root.label_10.text())

        document.add_paragraph(f"\nОбщее кол-во нир: {self.w_root.label_9.text()}")
        document.add_paragraph(f"Суммарный объем финансирования: {self.w_root.label_10.text()}")

        filename, _ = QFileDialog.getSaveFileName(
            None, 'Save Doc', os.getcwd())
        if filename:
            document.save(filename)

    def set_analyze_filters(self, filter_widget):
        if filter_widget:
            text = f"Федеральный округ: {filter_widget.w_root.comboBox_2.currentText()}\n" \
                   f"Субъект федерации: {filter_widget.w_root.comboBox_3.currentText()}\n" \
                   f"Город: {filter_widget.w_root.comboBox_4.currentText()}\n" \
                   f"ВУЗ: {filter_widget.w_root.comboBox_5.currentText()}\n" \
                   f"Первые цифры кода ГРНТИ: {filter_widget.w_root.lineEdit.text()}\n"

            self.w_root.textBrowser.setText(text)

    def set_anlyze_sum(self, last_query):
        query = QSqlQuery(f"""
                                    SELECT COUNT(*), SUM(f18) FROM ({last_query})
                                    """)
        while query.next():
            amount_of_nirs = query.value(0)
            sum_fin = query.value(1)

        self.w_root.label_9.setText(str(amount_of_nirs))
        self.w_root.label_10.setText(str(sum_fin))

    def open_analyze(self):
        self.set_analyze_filters(self.filter_nir_form)
        self.set_anlyze_sum(self.nir_model.query().lastQuery())
        self.anylyze_by_vuz.update(self.nir_model.query().lastQuery())
        self.anylyze_by_grnti.update(self.nir_model.query().lastQuery())
        self.anylyze_by_char.update(self.nir_model.query().lastQuery())

    def reset_filters(self):
        self.nir_model.reset_filters()
        self.filter_nir_form.reset()

    def open_filter_form(self):
        # self.filter_nir_form.reset()
        self.filter_nir_form.w.show()

    def add_record_nir_form(self):
        self.edit_record_nir_form.clear_form()
        self.edit_record_nir_form.w_root.pushButton_2.setText('Добавить запись')
        self.edit_record_nir_form.w_root.label_2.setText('Добавление записи')

        vuzes = self.vuz_model.get_list_vuzes_z2()

        ## ВУЗ
        self.edit_record_nir_form.w_root.comboBox_2.addItems(vuzes)

        self.edit_record_nir_form.w_root.pushButton_2.disconnect()
        self.edit_record_nir_form.w_root.pushButton_2.clicked.connect(lambda: self.change_add_record(add=True))

        self.edit_record_nir_form.w.show()

    def open_edit_record_nir(self):
        self.edit_record_nir_form.w_root.pushButton_2.setText('Принять')
        self.edit_record_nir_form.w_root.label_2.setText('Редактирование записи')
        self.edit_record_nir_form.clear_form()

        selected_rows = self.w_root.tableView.selectionModel().selectedRows()

        if len(selected_rows) == 1:
            edit_row = self.nir_model.row_from_index(selected_rows[0])
            vuzes = self.vuz_model.get_list_vuzes_z2()

            values_to_init_in_form = {
                "vuzes": vuzes,
                "current_vuz": edit_row['z2'],
                "reg_nomer": edit_row['rnw'],
                "type_of_nir": edit_row['f1'],
                "head_name": edit_row['f6'],
                "key_grnti": edit_row['f10'],
                "amount_of_funding": edit_row['f18'],
                "name_of_nir": edit_row['f2'],
                "head_position": edit_row['f7'],
            }

            self.edit_record_nir_form.set_vales(values_to_init_in_form)
            # ----------------------------------------------------------
            try:
                self.edit_record_nir_form.w_root.pushButton_2.clicked.disconnect()
            except:
                pass
            self.edit_record_nir_form.w_root.pushButton_2.clicked.connect(
                lambda: self.change_add_record(add=False, changed_codvuz=edit_row['codvuz'],
                                               changed_rnw=edit_row['rnw'], selected_row=selected_rows[0].row()))

            self.edit_record_nir_form.w.show()

    def change_add_record(self, add, changed_codvuz=0, changed_rnw='0', selected_row=None):

        edited_row = self.edit_record_nir_form.get_data()
        if not edited_row['error']:
            edited_row['codvuz'] = self.vuz_model.get_codvuz_from_z2(edited_row['z2'])
            if not edited_row['f18']:
                edited_row['f18'] = 0

            if not add:
                print(edited_row)
                self.nir_model.update_row(edited_row, changed_codvuz, changed_rnw)
                self.edit_record_nir_form.w.close()

            else:
                if not self.nir_model.add_row(edited_row):
                    self.edit_record_nir_form.w_root.label_11.setText(
                        'НИР с таким номером в данном вузе уже существует')
                else:
                    self.nir_model.reset_filters()
                    self.edit_record_nir_form.w.close()

            self.fin_vuz_model.recalculate_row([edited_row])
            row_to_select = self.nir_model.get_indexes_of_rows(edited_row)
            self.w_root.tableView.selectRow(row_to_select)
            self.w_root.tableView.scrollTo(self.nir_model.index(row_to_select, 0))

    def delete_records_nir(self):
        def accept_delete(rows_to_delete):
            # self.nir_model.delete_rows(rows_to_delete)
            self.nir_model.delete_rows(self.w_root.tableView.selectionModel().selectedRows())
            self.fin_vuz_model.recalculate_row(rows_to_delete)
            self.delete_accept_form.w.close()

        rows_to_delete = []
        for index in self.w_root.tableView.selectionModel().selectedRows():
            rows_to_delete.append(self.nir_model.row_from_index(index))
        if rows_to_delete:
            self.delete_accept_form.w.show()
            self.delete_accept_form.w_root.pushButton.clicked.connect(lambda: accept_delete(rows_to_delete))

    def onNirHeaderClicked(self, index):
        self.w_root.tableView.scrollTo(self.nir_model.index(0, 0))
        if self.nir_model.sortOrder[index] == 'ASC':
            self.nir_model.sortOrder[index] = 'DESC'
            self.w_root.tableView.horizontalHeader().setSortIndicator(index, Qt.SortOrder.AscendingOrder)
        else:
            self.nir_model.sortOrder[index] = 'ASC'
            self.w_root.tableView.horizontalHeader().setSortIndicator(index, Qt.SortOrder.DescendingOrder)

        if index != 0:
            self.w_root.tableView.horizontalHeader().setSortIndicatorShown(True)
            self.nir_model.orderByQuery = f""" ORDER BY {self.nir_model.record().fieldName(index) + ' ' + self.nir_model.sortOrder[index]} """
        else:
            self.nir_model.orderByQuery = f""" ORDER BY codvuz {self.nir_model.sortOrder[index]}, LENGTH(rnw), rnw """

        self.nir_model.update_model()

    def sort_by_codvuz_rnw(self):
        self.w_root.tableView.scrollTo(self.nir_model.index(0, 0))
        self.w_root.tableView.horizontalHeader().setSortIndicator(0, Qt.SortOrder.AscendingOrder)
        self.nir_model.orderByQuery = f""" ORDER BY codvuz, LENGTH(rnw), rnw """
        self.nir_model.update_model()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    ex = App()
    app.exec()
